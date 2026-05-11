from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


TITLE = "Atenxion Call Center User Manual"
SUBTITLE = "Realtime multi-agent telecom support demo"

OVERVIEW = [
    "The Atenxion Call Center Lab is a realtime telecom-support simulation built on the OpenAI Realtime Agents stack. It demonstrates specialist handoffs, local function-tool execution, output guardrails, and event-driven filler audio inside a transcript-first operator UI.",
    "The root scenario key is callcenteragent. The root agent greets, verifies, and routes callers into specialist flows for billing, technical support, retention, supervisor review, and human-style escalation.",
    "All account and policy answers are intentionally derived from deterministic local mock tools. This makes the environment predictable for demos, prompt tuning, and tool-call inspection without requiring a live backend.",
]

ARCHITECTURE_BULLETS = [
    "Scenario registry: callcenteragent is added alongside the existing simpleHandoff, customerServiceRetail, and chatSupervisor demos.",
    "Realtime agents: one shared voice is used across the Atenxion graph so handoffs remain stable during live audio sessions.",
    "Tool pipeline: specialists call local function tools for account lookup, billing interpretation, diagnostics, retention offers, and supervisor decisions.",
    "Guardrails: all Atenxion responses pass through the existing output moderation guardrail, now configured with Atenxion as the company context.",
    "UI surfacing: transcript breadcrumbs visually distinguish handoffs, tool starts, tool results, guardrail actions, and filler-audio markers.",
]

AGENTS = [
    (
        "callcenteragent",
        "Front-desk triage and verification",
        "Greeting, account verification, case setup, initial routing",
        "Routes to billing, technical support, retention, supervisor, or human escalation based on intent and caller state.",
    ),
    (
        "billingAgent",
        "Billing specialist",
        "Bill review, payment flexibility, goodwill credit handling",
        "Escalates to supervisorAgent for out-of-policy credits or sensitive exceptions.",
    ),
    (
        "technicalSupportAgent",
        "Technical support specialist",
        "Outages, diagnostics, device recovery, technician scheduling",
        "Hands to billing when service remediation needs financial follow-up, or to supervisorAgent for policy-sensitive edge cases.",
    ),
    (
        "retentionAgent",
        "Retention specialist",
        "Save offers, downgrade paths, cancellation risk handling",
        "Escalates to supervisorAgent for unsupported discounts and to humanEscalationAgent for emotionally difficult calls.",
    ),
    (
        "supervisorAgent",
        "Policy and exception authority",
        "Policy lookup, exception review, escalation decisions",
        "Can resolve difficult cases directly or pass to humanEscalationAgent when tone recovery is needed.",
    ),
    (
        "humanEscalationAgent",
        "Live-escalation style closer",
        "Calming the caller, summarizing history, closing the loop",
        "May return to supervisorAgent for final authority if required.",
    ),
]

TOOLS = [
    ("lookupCustomerProfile", "Shared", "Fetch the customer profile from the local mock account record."),
    ("verifyCaller", "Shared", "Validate phone number, DOB, and 4-digit PIN before account-specific help."),
    ("lookupActiveServices", "Shared", "List the lines and services on the account."),
    ("createCase", "Shared", "Open a support case with reason, priority, and team owner."),
    ("addCaseNote", "Shared", "Attach structured internal notes for downstream continuity."),
    ("getLatestBill", "Billing", "Retrieve the latest bill summary and due date."),
    ("explainChargeBreakdown", "Billing", "Translate the bill line items into plain-language explanations."),
    ("offerPaymentArrangement", "Billing", "Return a mock short-term payment plan offer."),
    ("applyGoodwillCredit", "Billing", "Apply or deny a goodwill credit based on authority limits."),
    ("checkServiceOutage", "Technical Support", "Check for an area outage by ZIP and service type."),
    ("runLineDiagnostics", "Technical Support", "Return mock line and device diagnostics."),
    ("scheduleTechnician", "Technical Support", "Schedule a technician appointment window."),
    ("rebootDeviceWorkflow", "Technical Support", "Trigger a remote reboot workflow for a registered device."),
    ("lookupPlanOptions", "Retention", "List Atenxion plan alternatives."),
    ("comparePlans", "Retention", "Compare the active plan to a target plan."),
    ("generateRetentionOffer", "Retention", "Return a mock save offer for churn-risk cases."),
    ("submitCancellationRequest", "Retention", "Submit a pending cancellation request."),
    ("lookupPolicyDocument", "Supervisor", "Search mock policy documents by topic."),
    ("approveException", "Supervisor", "Approve or deny a policy exception request."),
    ("escalationDecision", "Supervisor", "Return the recommended supervisor stance for a difficult call."),
]

UI_NOTES = [
    ("Scenario selector", "Switches the top-level agent graph between demos. callcenteragent is the new Atenxion flow."),
    ("Agent selector", "Lets you reconnect with a different Atenxion specialist as the root agent for testing."),
    ("Transcript", "Primary workspace for live messages plus structured breadcrumbs for handoffs, tools, audio, and guardrails."),
    ("Logs pane", "Transport-level and session event trace for deeper debugging."),
    ("Push to talk", "Toggles between server VAD and manual press-to-talk interaction."),
    ("Audio playback", "Enables or mutes remote assistant audio."),
    ("Codec control", "Switches between Opus, PCMU, and PCMA to preview different telephony fidelity modes."),
]

TEST_SCENARIOS = [
    (
        "Billing explanation",
        "Ask why the latest bill is high. Confirm that triage verifies the caller, routes to billingAgent, calls getLatestBill, and then calls explainChargeBreakdown.",
    ),
    (
        "Payment hardship",
        "State that you cannot pay the full balance this week. Billing should review the account and use offerPaymentArrangement.",
    ),
    (
        "Outage escalation",
        "Report that home internet is down in ZIP 98109. Technical support should checkServiceOutage and set outage expectations.",
    ),
    (
        "Technician dispatch",
        "Describe repeated home internet drops after reboot attempts. Technical support should runLineDiagnostics and scheduleTechnician.",
    ),
    (
        "Retention save path",
        "Say the service is getting too expensive and you may cancel. Retention should use lookupPlanOptions, comparePlans, and generateRetentionOffer.",
    ),
    (
        "Policy exception",
        "Request a credit larger than standard frontline authority. Billing should escalate to supervisorAgent, which uses lookupPolicyDocument and approveException.",
    ),
    (
        "Human escalation feel",
        "Act frustrated and ask for a real person. The flow should transfer to humanEscalationAgent and preserve context.",
    ),
]


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    if color:
      run.font.color.rgb = color


def configure_docx_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor(15, 23, 42)

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Aptos Display"
    heading1.font.size = Pt(16)
    heading1.font.color.rgb = RGBColor(30, 64, 175)

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Aptos"
    heading2.font.size = Pt(12)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(15, 23, 42)


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.2


def add_bullets(doc: Document, bullets: list[str]) -> None:
    for text in bullets:
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].width = Inches(widths[idx])
        shade_cell(header_cells[idx], "DCEAFE")
        set_cell_text(header_cells[idx], header, bold=True, color=RGBColor(30, 64, 175))

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            set_cell_text(cells[idx], value)

    doc.add_paragraph("")


def build_docx(path: Path) -> None:
    doc = Document()
    configure_docx_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(SUBTITLE)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(71, 85, 105)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_run = intro.add_run(
        "This guide covers the Atenxion telecom support scenario, its specialist agents, the mock tool surface, UI states, audio cues, and recommended test flows."
    )
    intro_run.font.size = Pt(11)
    intro_run.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_page_break()

    doc.add_heading("System Overview", level=1)
    add_paragraphs(doc, OVERVIEW)

    doc.add_heading("Scenario Architecture", level=1)
    add_bullets(doc, ARCHITECTURE_BULLETS)

    doc.add_heading("Agent Catalog", level=1)
    add_table(
        doc,
        ["Agent", "Role", "Primary Responsibilities", "Handoff Behavior"],
        AGENTS,
        [1.3, 1.35, 2.15, 1.7],
    )

    doc.add_heading("Tool Catalog", level=1)
    add_table(
        doc,
        ["Tool", "Group", "Purpose"],
        TOOLS,
        [1.7, 1.25, 3.2],
    )

    doc.add_heading("Guardrails and Filler Audio", level=1)
    add_paragraphs(
        doc,
        [
            "All Atenxion assistant responses continue to flow through the existing output moderation guardrail. In the UI, each assistant bubble surfaces a pass, pending, or flagged guardrail state, and any corrective action is visible in the transcript as a structured breadcrumb.",
            "Filler audio is event-driven only. Transfer ringing starts when an agent handoff occurs. Typing audio is used during tool-heavy wait moments. Both stop when the next assistant speech begins so the experience feels like live call-floor ambience rather than a permanent background loop.",
        ],
    )

    doc.add_heading("UI Walkthrough", level=1)
    add_table(
        doc,
        ["UI Element", "Behavior"],
        UI_NOTES,
        [1.7, 4.45],
    )

    doc.add_heading("Suggested End-to-End Tests", level=1)
    add_table(
        doc,
        ["Scenario", "What to Verify"],
        TEST_SCENARIOS,
        [1.55, 4.6],
    )

    conclusion = doc.add_paragraph()
    conclusion.paragraph_format.space_before = Pt(8)
    conclusion_run = conclusion.add_run(
        "For best validation, test the scenario with transcript breadcrumbs visible and the logs pane available for deeper inspection. The Atenxion demo is designed to be conversationally realistic while keeping tool behavior deterministic."
    )
    conclusion_run.italic = True

    doc.save(path)


def build_pdf(path: Path) -> None:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ManualTitle",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontName="Helvetica-Bold",
            fontSize=22,
            textColor=colors.HexColor("#0f172a"),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ManualSubtitle",
            parent=styles["Normal"],
            alignment=TA_CENTER,
            fontName="Helvetica",
            fontSize=11,
            textColor=colors.HexColor("#475569"),
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionHeading",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            textColor=colors.HexColor("#1e40af"),
            spaceBefore=10,
            spaceAfter=8,
        )
    )

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
    )

    story = [
        Paragraph(TITLE, styles["ManualTitle"]),
        Paragraph(SUBTITLE, styles["ManualSubtitle"]),
        Paragraph(
            "This guide covers the Atenxion telecom support scenario, its specialist agents, the mock tool surface, UI states, audio cues, and recommended test flows.",
            styles["ManualSubtitle"],
        ),
        Spacer(1, 0.15 * inch),
    ]

    def add_section(title: str, paragraphs: list[str]) -> None:
        story.append(Paragraph(title, styles["SectionHeading"]))
        for paragraph in paragraphs:
            story.append(Paragraph(paragraph, styles["BodyText"]))
            story.append(Spacer(1, 0.07 * inch))

    add_section("System Overview", OVERVIEW)
    add_section("Scenario Architecture", ARCHITECTURE_BULLETS)

    def add_pdf_table(headers: list[str], rows: list[tuple[str, ...]], widths: list[float]) -> None:
        data = [headers] + [list(row) for row in rows]
        table = Table(data, colWidths=[width * inch for width in widths], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                    ("LEADING", (0, 0), (-1, -1), 10),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 0.14 * inch))

    story.append(Paragraph("Agent Catalog", styles["SectionHeading"]))
    add_pdf_table(
        ["Agent", "Role", "Primary Responsibilities", "Handoff Behavior"],
        AGENTS,
        [1.15, 1.15, 2.0, 1.6],
    )

    story.append(PageBreak())
    story.append(Paragraph("Tool Catalog", styles["SectionHeading"]))
    add_pdf_table(["Tool", "Group", "Purpose"], TOOLS, [1.5, 1.1, 3.6])

    add_section(
        "Guardrails and Filler Audio",
        [
            "All Atenxion assistant responses continue to flow through the existing output moderation guardrail. In the UI, each assistant bubble surfaces a pass, pending, or flagged guardrail state, and any corrective action is visible in the transcript as a structured breadcrumb.",
            "Filler audio is event-driven only. Transfer ringing starts when an agent handoff occurs. Typing audio is used during tool-heavy wait moments. Both stop when the next assistant speech begins so the experience feels like live call-floor ambience rather than a permanent background loop.",
        ],
    )

    story.append(Paragraph("UI Walkthrough", styles["SectionHeading"]))
    add_pdf_table(["UI Element", "Behavior"], UI_NOTES, [1.4, 4.8])

    story.append(Paragraph("Suggested End-to-End Tests", styles["SectionHeading"]))
    add_pdf_table(["Scenario", "What to Verify"], TEST_SCENARIOS, [1.4, 4.8])

    doc.build(story)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    docs_dir = root / "docs"
    docs_dir.mkdir(exist_ok=True)

    docx_path = docs_dir / "atenxion-callcenter-user-manual.docx"
    pdf_path = docs_dir / "atenxion-callcenter-user-manual.pdf"

    build_docx(docx_path)
    build_pdf(pdf_path)

    print(docx_path)
    print(pdf_path)


if __name__ == "__main__":
    main()
